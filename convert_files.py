# from flask import Flask, request, send_file
# from werkzeug.utils import secure_filename
# import subprocess
# import os
# import uuid
# import img2pdf
# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# import io
# import logging

# # Configure logging
# logging.basicConfig(level=logging.DEBUG)
# logger = logging.getLogger(__name__)

# app = Flask(__name__)
# FILES_FOLDER = './files'  # Single directory for input & output

# # Ensure the directory exists
# os.makedirs(FILES_FOLDER, exist_ok=True)

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in {
#         'ppt', 'pptx', 'xls', 'xlsx', 'jpg', 'jpeg', 'png', 'gif', 'bmp', 'odt', 'rtf', 'txt', 'docx'
#     }

# def kill_libreoffice():
#     """Kill existing LibreOffice processes to prevent conflicts."""
#     try:
#         subprocess.run(["taskkill", "/IM", "soffice.bin", "/F"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
#         logger.info("Killed existing LibreOffice processes")
#     except subprocess.CalledProcessError:
#         logger.warning("No existing LibreOffice processes found.")

# def convert_docx_to_pdf(input_path, output_path):
#     """Convert .docx to PDF using ReportLab."""
#     try:
#         doc = Document(input_path)
#         pdf = canvas.Canvas(output_path, pagesize=letter)
#         y = 750
#         for para in doc.paragraphs:
#             if y < 50:
#                 pdf.showPage()
#                 y = 750
#             pdf.drawString(50, y, para.text)
#             y -= 12
#         pdf.save()
#         return True
#     except Exception as e:
#         logger.error(f"Error converting DOCX to PDF: {str(e)}", exc_info=True)
#         return False

# def convert_file(input_path, output_path):
#     """Convert a file to PDF based on its extension."""
#     try:
#         kill_libreoffice()  # Ensure no LibreOffice instance is running
#         input_path = os.path.abspath(input_path).replace("\\", "/")
#         output_path = os.path.abspath(output_path).replace("\\", "/")
        
#         if not os.path.exists(input_path):
#             logger.error(f"Error: Source file {input_path} not found!")
#             return False

#         file_ext = os.path.splitext(input_path)[1].lower()
        
#         if file_ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp'):
#             with open(input_path, "rb") as image_file:
#                 pdf_bytes = img2pdf.convert(image_file.read())
#             with open(output_path, "wb") as pdf_file:
#                 pdf_file.write(pdf_bytes)
#             return True
        
#         elif file_ext == '.docx':
#             return convert_docx_to_pdf(input_path, output_path)
        
#         elif file_ext in ('.ppt', '.pptx', '.xls', '.xlsx'):
#             command = [
#                 r'C:\Program Files\LibreOffice\program\soffice.exe',
#                 '--headless', '--norestore', '--nodefault', '--convert-to', 'pdf',
#                 '--outdir', FILES_FOLDER, input_path
#             ]
#             logger.debug(f"Executing: {' '.join(command)}")
#             result = subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
#             logger.debug(f"LibreOffice output:\n{result.stdout}")
#             logger.debug(f"LibreOffice errors:\n{result.stderr}")
#             return os.path.exists(output_path)
        
#         logger.warning(f"Unsupported file type: {file_ext}")
#         return False
#     except subprocess.CalledProcessError as e:
#         logger.error(f"LibreOffice error (code {e.returncode}): {e.stderr}")
#         return False
#     except Exception as e:
#         logger.error(f"Unexpected error during conversion: {str(e)}", exc_info=True)
#         return False

# @app.route('/convert', methods=['POST'])
# def handle_conversion():
#     if 'file' not in request.files:
#         return 'No file uploaded', 400
    
#     file = request.files['file']
#     if not file or file.filename == '' or not allowed_file(file.filename):
#         return 'Invalid file type', 400
    
#     try:
#         original_filename = secure_filename(file.filename)
#         unique_id = uuid.uuid4().hex
#         upload_path = os.path.join(FILES_FOLDER, f"{unique_id}_{original_filename}")
#         output_path = os.path.join(FILES_FOLDER, f"{unique_id}_converted.pdf")
        
#         file.save(upload_path)
#         logger.info(f"File saved to {upload_path}")
        
#         if not convert_file(upload_path, output_path):
#             return 'Conversion failed', 500
        
#         os.remove(upload_path)  # Cleanup original file
#         return send_file(output_path, mimetype='application/pdf')
    
#     except Exception as e:
#         logger.error(f"Server error: {str(e)}", exc_info=True)
#         return 'Internal server error', 500

# if __name__ == '__main__':
#     app.run(host='0.0.0.0', port=5001, debug=False)

import pdfplumber
from docx import Document

def pdf_to_word(pdf_path, docx_path):
    doc = Document()  # Create a new Word document

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)  # Add extracted text to Word file
    
    doc.save(docx_path)  # Save as DOCX
    print(f"Conversion successful: {docx_path}")

# Usage
pdf_to_word("./files/chumm.pdf", "./files/tttut.docx")
